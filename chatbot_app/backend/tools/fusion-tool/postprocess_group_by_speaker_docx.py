import argparse
import sys
import re
import html
from pathlib import Path
from typing import List, Optional, Tuple


try:
    from docx import Document
except Exception as e:  # pragma: no cover
    sys.stderr.write(
        "Missing dependency: python-docx. Install with: python -m pip install python-docx\n"
    )
    raise


TIMESTAMP_SPEAKER_RE = re.compile(r"^\[(?P<ts>\d{1,2}:\d{2}(?::\d{2})?)\]\s+(?P<speaker>.+?):\s*(?P<text>.*)\s*$")


def parse_line(line: str) -> Optional[Tuple[str, str, str]]:
    match = TIMESTAMP_SPEAKER_RE.match(line)
    if not match:
        return None
    ts = match.group("ts")
    speaker = match.group("speaker")
    text = match.group("text")
    return ts, speaker, text


def iter_groups(input_text: str, decode_html: bool):
    current_speaker: Optional[str] = None
    current_first_ts: Optional[str] = None
    current_items: List[Tuple[str, str]] = []

    def flush_group():
        nonlocal current_speaker, current_first_ts, current_items
        if current_speaker is not None and current_first_ts is not None and current_items:
            yield (current_speaker, current_first_ts, list(current_items))
        current_speaker = None
        current_first_ts = None
        current_items = []

    # We implement a generator pattern using a local buffer because Python doesn't allow yield in nested function easily
    output_buffer: List[Tuple[str, str, List[Tuple[str, str]]]] = []

    for raw_line in input_text.splitlines():
        line = raw_line.rstrip("\n\r")
        parsed = parse_line(line)
        if parsed is None:
            # Flush any open group; pass-through non-matching lines separately (handled by caller)
            for g in flush_group() or ():
                output_buffer.append(g)
            output_buffer.append(("__RAW__", "", [("", line)]))
            continue

        ts, speaker, text = parsed
        if decode_html:
            speaker = html.unescape(speaker)
            text = html.unescape(text)

        if current_speaker is None:
            current_speaker = speaker
            current_first_ts = ts
            current_items = [(ts, text)]
            continue

        if speaker == current_speaker:
            current_items.append((ts, text))
        else:
            for g in flush_group() or ():
                output_buffer.append(g)
            current_speaker = speaker
            current_first_ts = ts
            current_items = [(ts, text)]

    for g in flush_group() or ():
        output_buffer.append(g)

    for entry in output_buffer:
        yield entry


def write_docx(
    groups_iter,
    output_path: Path,
    merge: bool,
) -> None:
    doc = Document()

    for speaker, first_ts, items in groups_iter:
        if speaker == "__RAW__":
            text = items[0][1]
            doc.add_paragraph(text)
            continue

        header = f"[{first_ts}] {speaker}:"
        doc.add_paragraph(header)

        if merge:
            joined = " ".join(t.strip() for _, t in items if t.strip())
            if joined:
                doc.add_paragraph(joined)
            doc.add_paragraph("")
            continue

        for ts, text in items:
            text = text.strip()
            if not text:
                continue
            doc.add_paragraph(f"[{ts}] {text}")
        doc.add_paragraph("")

    # Remove a trailing blank paragraph if present is non-trivial; saving as-is
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(
        description=(
            "Export grouped transcript to Word (.docx). Mirrors the text grouper: "
            "default keeps per-line timestamps; --merge collapses to one paragraph per speaker."
        )
    )
    parser.add_argument(
        "--input",
        "-i",
        required=True,
        help="Path to input transcript (.txt or .docx)",
    )
    parser.add_argument(
        "--output",
        "-o",
        help=(
            "Path to output .docx. Defaults to '<input_without_ext>_grouped.docx' in the same folder."
        ),
    )
    parser.add_argument(
        "--no-decode-html",
        action="store_true",
        help="Do not HTML-unescape speaker names/text (default is to unescape).",
    )
    parser.add_argument(
        "--merge",
        action="store_true",
        help=(
            "Merge all lines of a speaker group into one paragraph (keeps only the first timestamp)."
        ),
    )

    args = parser.parse_args()
    input_path = Path(args.input)
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_name(input_path.stem + "_grouped.docx")

    # Support both .txt and .docx as input
    if input_path.suffix.lower() == ".docx":
        doc = Document(str(input_path))
        text = "\n".join(p.text for p in doc.paragraphs)
    else:
        text = input_path.read_text(encoding="utf-8")
    groups = iter_groups(
        input_text=text,
        decode_html=not args.no_decode_html,
    )
    write_docx(groups_iter=groups, output_path=output_path, merge=args.merge)


if __name__ == "__main__":
    main()


