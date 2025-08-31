from pathlib import Path
from typing import Dict, List
import re

from docx import Document

# Pattern A (rare): [mm:ss] Speaker: Text
_BRACKET_LINE_RE = re.compile(r"^\[(\d{2}):(\d{2})\]\s*(.+?):\s*(.*)$")
# Pattern B (Teams export seen in screenshots): 'Speaker   0:11' header line,
# followed by one or more content lines, until the next header.
_HEADER_RE = re.compile(r"^(?P<speaker>.+?)\s+(?:(?P<h>\d{1,2})[:\uFF1A\u2236])?(?P<m>\d{1,2})[:\uFF1A\u2236](?P<s>\d{2})$")
_TIME_ONLY_RE = re.compile(r"^(?:(?P<h>\d{1,2})[:\uFF1A\u2236])?(?P<m>\d{1,2})[:\uFF1A\u2236](?P<s>\d{2})$")
_ANY_TIME_RE = re.compile(r"(?:(?P<h>\d{1,2})[:\uFF1A\u2236])?(?P<m>\d{1,2})[:\uFF1A\u2236](?P<s>\d{2})")


def _to_seconds(h: int, m: int, s: int) -> int:
    return h * 3600 + m * 60 + s


def _mmss_to_seconds(mm: str, ss: str) -> int:
    return int(mm) * 60 + int(ss)


def _yield_candidate_lines(doc: Document):
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            yield text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                if not cell_text:
                    continue
                for line in cell_text.splitlines():
                    line = line.strip()
                    if line:
                        yield line


_META_PREFIXES = {
    "transkript",
}
_META_CONTAINS = {
    "transkription gestartet",
    "transkription beendet",
}


def _is_meta(line: str) -> bool:
    lower = line.lower()
    if any(lower.startswith(p) for p in _META_PREFIXES):
        return True
    if any(s in lower for s in _META_CONTAINS):
        return True
    # dates like '20. August 2025, 07:03AM'
    if "," in line and any(m in lower for m in ["januar", "februar", "märz", "april", "mai", "juni", "juli", "august", "september", "oktober", "november", "dezember"]):
        return True
    return False


def read_teams_docx(path: Path) -> List[Dict]:
    segments: List[Dict] = []
    doc = Document(str(path))

    # First: parse any bracket-style lines if present
    for line in _yield_candidate_lines(doc):
        line = line.replace("\u00A0", " ").replace("\uFF1A", ":").replace("\u2236", ":")
        if _is_meta(line):
            continue
        m = _BRACKET_LINE_RE.match(line)
        if not m:
            continue
        mm, ss, speaker, content = m.groups()
        t = _mmss_to_seconds(mm, ss)
        segments.append({
            "source": "teams",
            "t_start": t,
            "t_end": None,
            "speaker": speaker.strip(),
            "text": content.strip(),
            "tokens": None,
        })
    if segments:
        return segments

    # Fallback: header + content blocks
    current = None
    buffer: List[str] = []
    prev_speaker: str | None = None
    for line in _yield_candidate_lines(doc):
        line = line.replace("\u00A0", " ").replace("\uFF1A", ":").replace("\u2236", ":")
        if _is_meta(line):
            continue
        m2 = _HEADER_RE.match(line)
        if m2:
            # flush previous
            if current is not None:
                segments.append({
                    "source": "teams",
                    "t_start": current["t_start"],
                    "t_end": None,
                    "speaker": current["speaker"],
                    "text": " ".join(x.strip() for x in buffer if x.strip()),
                    "tokens": None,
                })
            buffer = []
            h = int(m2.group("h") or 0)
            m = int(m2.group("m"))
            s = int(m2.group("s") or 0)
            speaker = m2.group("speaker").strip()
            current = {"t_start": _to_seconds(h, m, s), "speaker": speaker}
            prev_speaker = speaker
            continue

        # time-only line on its own, previous paragraph likely speaker name
        t_only = _TIME_ONLY_RE.match(line)
        if t_only and prev_speaker:
            if current is not None:
                segments.append({
                    "source": "teams",
                    "t_start": current["t_start"],
                    "t_end": None,
                    "speaker": current["speaker"],
                    "text": " ".join(x.strip() for x in buffer if x.strip()),
                    "tokens": None,
                })
            buffer = []
            h = int(t_only.group("h") or 0)
            m = int(t_only.group("m"))
            s = int(t_only.group("s") or 0)
            current = {"t_start": _to_seconds(h, m, s), "speaker": prev_speaker}
            continue

        # Plain speaker line (no colon/pipe) – remember as potential prev speaker
        if ":" not in line and "|" not in line and len(line.split()) <= 4:
            prev_speaker = line.strip()
            continue

        # Heuristic: time appears anywhere in the line (e.g., 'Bastian   0:11' with extra text)
        anyt = _ANY_TIME_RE.search(line)
        if anyt:
            if current is not None:
                segments.append({
                    "source": "teams",
                    "t_start": current["t_start"],
                    "t_end": None,
                    "speaker": current["speaker"],
                    "text": " ".join(x.strip() for x in buffer if x.strip()),
                    "tokens": None,
                })
            buffer = []
            h = int(anyt.group("h") or 0)
            m = int(anyt.group("m"))
            s = int(anyt.group("s") or 0)
            speaker = prev_speaker or line[:anyt.start()].strip()
            current = {"t_start": _to_seconds(h, m, s), "speaker": speaker}
            continue

        buffer.append(line)
    if current is not None:
        segments.append({
            "source": "teams",
            "t_start": current["t_start"],
            "t_end": None,
            "speaker": current["speaker"],
            "text": " ".join(x.strip() for x in buffer if x.strip()),
            "tokens": None,
        })
    return segments
