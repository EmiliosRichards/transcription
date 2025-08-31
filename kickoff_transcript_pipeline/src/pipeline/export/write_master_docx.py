from pathlib import Path
from typing import List, Dict, Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _format_mmss(seconds: int) -> str:
    m = int(seconds) // 60
    s = int(seconds) % 60
    return f"{m:02d}:{s:02d}"


def _add_toc(document: Document) -> None:
    # Insert a TOC field: Word will populate on open (References > Update Table)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fld_char_begin)
    r.append(instr_text)
    r.append(fld_char_separate)
    r.append(fld_char_end)


def write_master_docx(lines: List[Dict[str, Any]], out_path: Path, chapter_minutes: int) -> None:
    document = Document()

    document.add_heading("Master Transcript", level=0)
    document.add_paragraph("Table of Contents (update on open)")
    _add_toc(document)

    # Add content with chapter headings every chapter_minutes
    current_chapter = None
    for line in lines:
        t = int(line.get("t") or line.get("t_start") or 0)
        chapter_index = t // (chapter_minutes * 60) if chapter_minutes and chapter_minutes > 0 else 0
        if current_chapter is None or chapter_index != current_chapter:
            start_sec = chapter_index * chapter_minutes * 60
            end_sec = start_sec + (chapter_minutes * 60) - 1
            heading = f"{_format_mmss(start_sec)}–{_format_mmss(end_sec)}"
            document.add_heading(heading, level=1)
            current_chapter = chapter_index

        speaker = str(line.get("speaker", "")).strip()
        text = str(line.get("text", "")).strip()
        timestamp = _format_mmss(t)
        paragraph_text = f"[{timestamp}] {speaker}: {text}" if speaker else f"[{timestamp}] {text}"
        document.add_paragraph(paragraph_text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
